from __future__ import annotations

import asyncio
import datetime
from io import BytesIO

import plotly.graph_objects as go

from .utils import sanitize_filename

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:  # pragma: no cover
    Document = None
    Inches = None


def _plotly_fig_dict_to_png_bytes(fig_dict, *, width=1200, height=700, scale=2) -> bytes | None:
    if not fig_dict:
        return None
    try:
        fig = go.Figure(fig_dict)
        return fig.to_image(format="png", width=width, height=height, scale=scale)
    except asyncio.CancelledError as e:
        # Python 3.12: CancelledError is BaseException, handle explicitly
        print(f"[REPORT] plotly to_image cancelled (kaleido): {e}")
        return None
    except Exception as e:
        print(f"[REPORT] plotly to_image failed (kaleido?): {e}")
        return None


def build_report_docx_bytes(report_data: dict, track_fig_dict, brake_fig_dict, accel_fig_dict) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx не установлен. Установи: pip install python-docx")

    now_str = datetime.datetime.now().strftime("%d.%m.%Y %H:%M:%S")

    doc = Document()
    doc.add_heading("F1 Telemetry Report", level=0)
    doc.add_paragraph(f"Сформировано: {now_str}")

    doc.add_heading("Параметры", level=1)
    doc.add_paragraph(
        f"Год: {report_data.get('year')}\n"
        f"Тип: {report_data.get('event_type')}\n"
        f"Событие: {report_data.get('event')}\n"
        f"Сессия: {report_data.get('session')}\n"
        f"Гонщик: {report_data.get('driver_name')} ({report_data.get('driver_code')})\n"
        f"Команда: {report_data.get('team')}\n"
        f"Круг: {report_data.get('lap_num')}"
    )

    doc.add_heading("Статистика круга", level=1)
    table = doc.add_table(rows=0, cols=2)

    def add_row(k, v):
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    add_row("Время круга", report_data.get("lap_time"))
    add_row("Sector 1", report_data.get("sector1"))
    add_row("Sector 2", report_data.get("sector2"))
    add_row("Sector 3", report_data.get("sector3"))
    add_row("Шины", report_data.get("compound"))
    add_row("Износ шин (круги)", report_data.get("tyre_life"))
    add_row("Пит-стопы до круга", report_data.get("pit_stops_before"))

    portrait_bytes = data_uri_to_bytes(report_data.get("portrait_src"))
    logo_bytes = data_uri_to_bytes(report_data.get("team_logo_src"))

    if portrait_bytes or logo_bytes:
        doc.add_heading("Изображения", level=1)

    if portrait_bytes:
        doc.add_paragraph("Портрет гонщика")
        doc.add_picture(BytesIO(portrait_bytes), width=Inches(1.6))

    if logo_bytes:
        doc.add_paragraph("Логотип команды")
        doc.add_picture(BytesIO(logo_bytes), width=Inches(1.6))

    doc.add_heading("Графики", level=1)

    track_png = _plotly_fig_dict_to_png_bytes(track_fig_dict)
    if track_png:
        doc.add_paragraph("Карта трассы")
        doc.add_picture(BytesIO(track_png), width=Inches(6.5))
    else:
        doc.add_paragraph("Карта трассы: не удалось выгрузить (проверь kaleido).")

    brake_png = _plotly_fig_dict_to_png_bytes(brake_fig_dict, width=1400, height=700, scale=2)
    if brake_png:
        doc.add_paragraph("Телеметрия (газ/тормоз/скорость)")
        doc.add_picture(BytesIO(brake_png), width=Inches(6.5))
    else:
        doc.add_paragraph("Телеметрия: не удалось выгрузить (проверь kaleido).")

    accel_png = _plotly_fig_dict_to_png_bytes(accel_fig_dict, width=1400, height=700, scale=2)
    if accel_png:
        doc.add_paragraph("Карта ускорений/торможений")
        doc.add_picture(BytesIO(accel_png), width=Inches(6.5))
    else:
        doc.add_paragraph("Карта ускорений: нет данных/не удалось выгрузить (проверь kaleido).")

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def default_report_filename(report_data: dict) -> str:
    fn = (
        f"F1_{report_data.get('year')}_{report_data.get('event')}_{report_data.get('session')}_"
        f"{report_data.get('driver_code')}_Lap{report_data.get('lap_num')}.docx"
    )
    return sanitize_filename(fn)

